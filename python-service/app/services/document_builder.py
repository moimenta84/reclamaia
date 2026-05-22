import os
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT


def build_documents(claim_id: int, text: str, output_dir: str) -> dict:
    """Generate Word (.docx) and PDF from Markdown claim text. Returns paths."""
    out = Path(output_dir) / "documents" / str(claim_id)
    out.mkdir(parents=True, exist_ok=True)

    word_path = str(out / f"reclamacion_{claim_id}.docx")
    pdf_path = str(out / f"reclamacion_{claim_id}.pdf")

    _build_docx(text, word_path)
    _build_pdf(text, pdf_path)

    # Return relative paths (from storage root)
    rel_word = f"documents/{claim_id}/reclamacion_{claim_id}.docx"
    rel_pdf = f"documents/{claim_id}/reclamacion_{claim_id}.pdf"

    return {
        "word_path": rel_word,
        "pdf_path": rel_pdf,
    }


def _clean_markdown(text: str) -> str:
    """Strip basic Markdown formatting for plain text insertion."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    return text


def _build_docx(text: str, path: str) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    title = doc.add_heading("RECLAMACIÓN FORMAL", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.add_paragraph()

    # Body — split by lines and render
    for line in text.split("\n"):
        clean = _clean_markdown(line).strip()
        if not clean:
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            p = doc.add_heading(clean, level=1)
        elif line.startswith("## "):
            p = doc.add_heading(clean, level=2)
        elif line.startswith("### "):
            p = doc.add_heading(clean, level=3)
        else:
            p = doc.add_paragraph(clean)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)

    # Legal disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "AVISO LEGAL: Este documento es orientativo. "
        "Consulte con un abogado para casos complejos. "
        "Generado por ReclamaIA — reclamaia.es"
    )
    disclaimer.runs[0].font.size = Pt(8)
    disclaimer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(path)


def _build_pdf(text: str, path: str) -> None:
    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        leftMargin=3 * cm,
        rightMargin=3 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=11,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=8,
    )
    heading1_style = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontSize=14,
        leading=18,
        spaceBefore=12,
        spaceAfter=6,
        alignment=TA_LEFT,
    )
    heading2_style = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=12,
        leading=16,
        spaceBefore=10,
        spaceAfter=4,
    )
    disclaimer_style = ParagraphStyle(
        "Disclaimer",
        parent=styles["Normal"],
        fontSize=8,
        textColor=(0.5, 0.5, 0.5),
        alignment=TA_LEFT,
        spaceBefore=20,
    )

    story = []

    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontSize=16,
        leading=20,
        spaceAfter=16,
        alignment=TA_LEFT,
    )
    story.append(Paragraph("RECLAMACIÓN FORMAL", title_style))
    story.append(Spacer(1, 0.5 * cm))

    for line in text.split("\n"):
        clean = _clean_markdown(line).strip()
        if not clean:
            story.append(Spacer(1, 0.3 * cm))
            continue

        if line.startswith("# "):
            story.append(Paragraph(clean, heading1_style))
        elif line.startswith("## "):
            story.append(Paragraph(clean, heading2_style))
        elif line.startswith("### "):
            story.append(Paragraph(clean, heading2_style))
        else:
            safe = clean.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe, body_style))

    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph(
        "AVISO LEGAL: Este documento es orientativo. "
        "Consulte con un abogado para casos complejos. "
        "Generado por ReclamaIA — reclamaia.es",
        disclaimer_style,
    ))

    doc.build(story)
